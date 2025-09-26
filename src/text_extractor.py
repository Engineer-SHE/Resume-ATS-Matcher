import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import PyPDF2
import pdfplumber
from docx import Document
import logging

logger = logging.getLogger(__name__)

class TextExtractor:
    def __init__(self):
        self.section_patterns = {
            'summary': r'(?i)(summary|objective|profile|about)',
            'experience': r'(?i)(experience|employment|work\s+history|professional\s+experience)',
            'education': r'(?i)(education|academic|qualification)',
            'skills': r'(?i)(skills|technical\s+skills|competenc)',
            'achievements': r'(?i)(achievement|accomplishment|award|certification)',
            'projects': r'(?i)(project|portfolio)'
        }

    def extract_from_file(self, file_path: str) -> str:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif extension == '.txt':
            return file_path.read_text(encoding='utf-8')
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def _extract_from_pdf(self, file_path: Path) -> str:
        text = ""

        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"

        return self._clean_text(text)

    def _extract_from_docx(self, file_path: Path) -> str:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text

        return self._clean_text(text)

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\@\#\$\%\&\*\(\)\[\]\{\}\/\\\|\'\"]', ' ', text)
        text = text.strip()
        return text

    def extract_sections(self, text: str) -> Dict[str, str]:
        sections = {
            'full_text': text,
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'achievements': '',
            'projects': ''
        }

        lines = text.split('\n')
        current_section = None
        section_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            section_found = False
            for section_name, pattern in self.section_patterns.items():
                if re.search(pattern, line):
                    if current_section and section_content:
                        sections[current_section] = '\n'.join(section_content).strip()
                    current_section = section_name
                    section_content = []
                    section_found = True
                    break

            if not section_found and current_section:
                section_content.append(line)
            elif not section_found and not current_section:
                if 'summary' not in sections or not sections['summary']:
                    sections['summary'] += line + ' '

        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content).strip()

        return sections

    def extract_contact_info(self, text: str) -> Dict[str, str]:
        contact = {}

        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{3,5}[-\s\.]?[0-9]{3,5}'
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin:)[\w\-]+'

        email_match = re.search(email_pattern, text)
        if email_match:
            contact['email'] = email_match.group()

        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact['phone'] = phone_match.group()

        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group()

        return contact